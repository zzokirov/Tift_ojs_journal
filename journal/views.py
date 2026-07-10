from django.db.models import Q
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth import login, update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import JournalIssue, Article, StaffMember, SiteVisit, Conference, News, Document
from .forms import ArticleSubmissionForm, CustomUserCreationForm, ProfileUpdateForm, CustomPasswordChangeForm


def index(request):
    # Login qilgan user bosh sahifaga kelsa — hisobiga yo'naltir
    if request.user.is_authenticated:
        return redirect('my_articles')

    # Tashrif buyuruvchilarni sanash
    try:
        ip = request.META.get('HTTP_X_FORWARDED_FOR', request.META.get('REMOTE_ADDR', '127.0.0.1'))
        if ',' in ip:
            ip = ip.split(',')[0].strip()
        from datetime import date, timedelta
        SiteVisit.objects.get_or_create(date=date.today(), ip_address=ip)
        total_visitors = SiteVisit.objects.count()
        today_visitors = SiteVisit.objects.filter(date=date.today()).count()
        # So'nggi 7 kunlik statistika
        week_labels = []
        week_data = []
        for i in range(6, -1, -1):
            d = date.today() - timedelta(days=i)
            week_labels.append(d.strftime('%d.%m'))
            week_data.append(SiteVisit.objects.filter(date=d).count())
    except Exception:
        total_visitors = 0
        today_visitors = 0
        week_labels = []
        week_data = []

    query = request.GET.get('q')
    recent_articles = Article.objects.filter(status='published').order_by('-created_at')
    if query:
        recent_articles = recent_articles.filter(
            Q(title__icontains=query) |
            Q(abstract__icontains=query) |
            Q(keywords__icontains=query) |
            Q(authors__icontains=query) |
            Q(author__first_name__icontains=query) |
            Q(author__last_name__icontains=query) |
            Q(author__username__icontains=query)
        )
    recent_articles = recent_articles[:10]
    issues = JournalIssue.objects.all().order_by('-year', '-number')
    return render(request, 'index.html', {
        'recent_articles': recent_articles,
        'issues': issues,
        'query': query,
        'total_visitors': total_visitors,
        'today_visitors': today_visitors,
        'week_labels': week_labels,
        'week_data': week_data,
    })


def archive(request):
    issues = JournalIssue.objects.all().order_by('-year', '-number')
    return render(request, 'archive.html', {'issues': issues})


def about(request):
    staff = StaffMember.objects.filter(is_active=True).order_by('order', 'full_name')
    areas = [
        "Arxitektura nazariyasi", "Binolar konstruksiyasi",
        "Shaharsozlik va landshaft", "Geodeziya va kartografiya",
        "Ta'lim metodikasi", "Raqamli texnologiyalar",
        "Sun'iy idrok", "Kadastr va er resurslari",
    ]
    return render(request, 'about.html', {'staff': staff, 'areas': areas})


def issue_detail(request, issue_pk):
    issue = get_object_or_404(JournalIssue, pk=issue_pk)
    articles = Article.objects.filter(issue=issue, status='published')
    return render(request, 'journal/issue_detail.html', {
        'issue': issue,
        'articles': articles,
    })


def article_detail(request, pk):
    article = get_object_or_404(Article, pk=pk)
    if not request.session.get(f'viewed_article_{pk}'):
        article.views_count += 1
        article.save()
        request.session[f'viewed_article_{pk}'] = True

    # Shu muallifning boshqa nashr etilgan maqolalari
    author_articles = Article.objects.filter(
        author=article.author,
        status='published'
    ).exclude(pk=pk).order_by('-created_at')[:8]

    return render(request, 'article_detail.html', {
        'article': article,
        'author_articles': author_articles,
    })


def _extract_pdf_text_as_html(pdf_path):
    """
    PDF fayldan matn, jadval va rasmlarni o'qib HTML qaytaradi.
    - Matn bloklari: paragraf/sarlavha sifatida
    - Rasm bloklari: PyMuPDF bilan render qilib base64 PNG sifatida
    PyMuPDF (fitz) ishlatiladi.
    """
    try:
        import fitz
        import html as html_lib
        import base64

        doc = fitz.open(pdf_path)
        result_html = []

        for page_idx, page in enumerate(doc):
            page_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            page_rect = page.rect

            for block in page_dict.get("blocks", []):
                block_type = block.get("type", -1)

                # ── RASM BLOKI (type=1) ──
                if block_type == 1:
                    try:
                        # Rasm bbox ni olish
                        bbox = block.get("bbox")
                        if not bbox:
                            continue
                        rect = fitz.Rect(bbox)
                        # Clip rect orqali faqat shu qismni render qilamiz
                        mat = fitz.Matrix(2.0, 2.0)  # 2x zoom — sifat uchun
                        clip = rect
                        pix = page.get_pixmap(matrix=mat, clip=clip)
                        img_bytes = pix.tobytes("png")
                        b64 = base64.b64encode(img_bytes).decode('ascii')
                        # Kenglikni pt da hisoblash (xhtml2pdf % ni tushunmaydi)
                        img_width_pt = rect.width
                        page_width_pt = page_rect.width
                        # Maksimal kenglik = sahifa kengligi (margins hisobga olingan)
                        content_width_pt = page_width_pt - 142  # ~5cm margin jami
                        render_width_pt = min(img_width_pt, content_width_pt)
                        result_html.append(
                            f'<div class="doc-img-wrap">'
                            f'<img src="data:image/png;base64,{b64}" '
                            f'style="width:{render_width_pt:.0f}pt;max-width:100%;height:auto;display:block;margin:6pt auto;"/>'
                            f'</div>'
                        )
                    except Exception:
                        pass
                    continue

                # ── MATN BLOKI (type=0) ──
                if block_type != 0:
                    continue

                block_lines = []
                for line in block.get("lines", []):
                    line_text = ""
                    for span in line.get("spans", []):
                        line_text += span.get("text", "")
                    stripped = line_text.strip()
                    if stripped:
                        block_lines.append(stripped)

                if not block_lines:
                    continue

                block_text = " ".join(block_lines)

                # Font o'lcham va bold tekshirish
                font_size = 11
                is_bold = False
                if block.get("lines"):
                    first_spans = block["lines"][0].get("spans", [])
                    if first_spans:
                        font_size = first_spans[0].get("size", 11)
                        flags = first_spans[0].get("flags", 0)
                        is_bold = bool(flags & 16)

                escaped = html_lib.escape(block_text)

                if font_size >= 14 or (font_size >= 12 and is_bold):
                    result_html.append(f'<h2 class="doc-heading">{escaped}</h2>')
                elif font_size >= 12:
                    result_html.append(f'<h3 class="doc-subheading">{escaped}</h3>')
                else:
                    result_html.append(f'<p>{escaped}</p>')

        doc.close()
        return '\n'.join(result_html)
    except Exception:
        return ''


def _extract_docx_text_as_html(docx_path):
    """
    Word (.docx) fayldan matn, sarlavha, jadval va rasmlarni HTML ga o'giradi.
    - Paragraflar: stil bo'yicha sarlavha/paragraf
    - Jadvallar: to'liq HTML jadval
    - Rasmlar: base64 PNG sifatida inline embed
    python-docx ishlatiladi.
    """
    try:
        import docx
        import html as html_lib
        import base64
        from docx.oxml.ns import qn
        from lxml import etree

        doc = docx.Document(docx_path)

        # Barcha relationships (rImage) dan rasm ma'lumotlarini olish
        # doc.part.rels: {rId: rel}
        def get_image_base64(rel_id):
            try:
                rel = doc.part.rels.get(rel_id)
                if rel and "image" in rel.reltype:
                    img_data = rel.target_part.blob
                    ext = rel.target_part.content_type.split('/')[-1]
                    if ext == 'jpeg':
                        ext = 'jpg'
                    b64 = base64.b64encode(img_data).decode('ascii')
                    return f"data:image/{ext};base64,{b64}"
            except Exception:
                pass
            return None

        result_html = []

        # Body ichidagi barcha elementlarni tartib bilan o'tamiz
        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            # ── PARAGRAF ──
            if tag == 'p':
                para = None
                for p in doc.paragraphs:
                    if p._element is child:
                        para = p
                        break
                if para is None:
                    continue

                # Paragrafda rasm borligini tekshirish (drawing/blipFill)
                drawings = child.findall('.//' + qn('a:blip'))
                if drawings:
                    # Rasmlar embed yoki link orqali
                    for blip in drawings:
                        r_embed = blip.get(qn('r:embed'))
                        r_link = blip.get(qn('r:link'))
                        rid = r_embed or r_link
                        if rid:
                            src = get_image_base64(rid)
                            if src:
                                result_html.append(
                                    f'<div class="doc-img-wrap">'
                                    f'<img src="{src}" style="max-width:100%;height:auto;display:block;margin:6pt auto;"/>'
                                    f'</div>'
                                )
                    # Agar paragrafda matn ham bor bo'lsa
                    text = para.text.strip()
                    if text:
                        escaped = html_lib.escape(text)
                        result_html.append(f'<p class="img-caption">{escaped}</p>')
                    continue

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name.lower() if para.style else ''
                escaped = html_lib.escape(text)

                if 'heading 1' in style_name or style_name == 'title':
                    result_html.append(f'<h2 class="doc-heading">{escaped}</h2>')
                elif 'heading 2' in style_name:
                    result_html.append(f'<h3 class="doc-subheading">{escaped}</h3>')
                elif 'heading' in style_name:
                    result_html.append(f'<h4 class="doc-subheading">{escaped}</h4>')
                else:
                    is_bold = any(run.bold for run in para.runs if run.text.strip())
                    if is_bold and len(text) < 120:
                        result_html.append(f'<p><strong>{escaped}</strong></p>')
                    else:
                        result_html.append(f'<p>{escaped}</p>')

            # ── JADVAL ──
            elif tag == 'tbl':
                tbl = None
                for t in doc.tables:
                    if t._element is child:
                        tbl = t
                        break
                if tbl is None:
                    continue

                table_html = ['<table class="doc-table">']
                for i, row in enumerate(tbl.rows):
                    table_html.append('<tr>')
                    for cell in row.cells:
                        cell_text = html_lib.escape(cell.text.strip())
                        if i == 0:
                            table_html.append(f'<th>{cell_text}</th>')
                        else:
                            table_html.append(f'<td>{cell_text}</td>')
                    table_html.append('</tr>')
                table_html.append('</table>')
                result_html.append('\n'.join(table_html))

        return '\n'.join(result_html)
    except Exception:
        return ''


def _get_file_as_bytes(article_file):
    """
    Fayl maydonidan bytes qaytaradi.
    Lokal (path) bo'lsa — disk dan o'qiydi.
    Cloudinary (URL) bo'lsa — HTTP orqali yuklab oladi.
    """
    import os
    try:
        # Lokal fayl
        path = article_file.path
        if os.path.exists(path):
            with open(path, 'rb') as f:
                return f.read(), path
    except (NotImplementedError, AttributeError, Exception):
        pass

    # Cloudinary yoki tashqi URL
    try:
        import requests as req
        url = article_file.url
        if url:
            resp = req.get(url, timeout=30)
            if resp.status_code == 200:
                # Vaqtinchalik fayl sifatida saqlash
                import tempfile
                suffix = '.' + url.split('?')[0].rsplit('.', 1)[-1]
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                tmp.write(resp.content)
                tmp.close()
                return resp.content, tmp.name
    except Exception:
        pass

    return None, None


def _get_article_content_html(article_file):
    """
    Fayl ob'ektidan (FileField) matn/rasm HTML qaytaradi.
    Lokal va Cloudinary (URL) ni qo'llab-quvvatlaydi.
    """
    if not article_file:
        return ''

    file_bytes, file_path = _get_file_as_bytes(article_file)
    if not file_path:
        return ''

    try:
        url_or_name = getattr(article_file, 'name', '') or ''
        ext = url_or_name.lower().rsplit('.', 1)[-1].split('?')[0]
        if ext in ('doc', 'docx'):
            return _extract_docx_text_as_html(file_path)
        else:
            return _extract_pdf_text_as_html(file_path)
    finally:
        # Vaqtinchalik faylni o'chirish
        import os, tempfile
        try:
            tmp_dir = tempfile.gettempdir()
            if file_path.startswith(tmp_dir):
                os.unlink(file_path)
        except Exception:
            pass


def _add_header_footer_to_pdf(pdf_bytes, article):
    """
    PyMuPDF yordamida asl PDF faylning har sahifasiga
    yuqori va pastki kolontitullarni qo'shadi.
    Matn, rasmlar, jadvallar o'zgarmaydi.
    """
    try:
        import fitz
        import io

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        # Ranglar
        DARK_BLUE = (0.04, 0.086, 0.157)   # #0a1628
        GREEN     = (0.086, 0.502, 0.239)   # #15803d
        GRAY      = (0.4, 0.4, 0.4)

        # Jurnal ma'lumotlari
        journal_name  = "TIFT JOURNAL"
        journal_sub   = '"Arxitektura va Ta\'lim" Ilmiy-Elektron Jurnali'
        footer_text   = "Toshkent sh., Amir Temur ko'chasi, 108  |  Tel: +998 71 238-74-80  |  journal@tift.uz  |  www.tift.uz"

        if article.issue:
            issue_text = f"Jild {article.issue.volume}, Son {article.issue.number}, {article.issue.year}"
        else:
            issue_text = "ISSN: 2181-XXXX"

        for page in doc:
            w = page.rect.width    # sahifa kengligi (pt)
            h = page.rect.height   # sahifa balandligi (pt)

            margin_x = 71.0  # 2.5cm
            header_y = 22.0  # yuqoridan 0.77cm
            footer_y = h - 25.0  # pastdan

            # ── YUQORI KOLONTITUL ──

            # Logo kvadrat (chap)
            logo_rect = fitz.Rect(margin_x, header_y - 2, margin_x + 28, header_y + 26)
            page.draw_rect(logo_rect, color=DARK_BLUE, fill=DARK_BLUE, width=0)
            page.insert_text(
                (margin_x + 7, header_y + 19),
                "T", fontsize=16,
                color=(0.133, 0.773, 0.369),  # #22c55e
                fontname="Helvetica-Bold"
            )

            # Jurnal nomi (o'rta)
            page.insert_text(
                (margin_x + 34, header_y + 10),
                journal_name, fontsize=10,
                color=DARK_BLUE, fontname="Helvetica-Bold"
            )
            page.insert_text(
                (margin_x + 34, header_y + 22),
                journal_sub, fontsize=7,
                color=GRAY, fontname="Helvetica"
            )

            # Jurnal soni (o'ng)
            page.insert_text(
                (w - margin_x - 100, header_y + 10),
                issue_text, fontsize=8,
                color=GRAY, fontname="Helvetica"
            )
            page.insert_text(
                (w - margin_x - 70, header_y + 20),
                "ISSN: 2181-XXXX", fontsize=7,
                color=GRAY, fontname="Helvetica"
            )

            # Yuqori chiziq (yashil)
            page.draw_line(
                (margin_x, header_y + 30),
                (w - margin_x, header_y + 30),
                color=GREEN, width=1.2
            )

            # ── PASTKI KOLONTITUL ──

            # Pastki chiziq (yashil)
            page.draw_line(
                (margin_x, footer_y - 4),
                (w - margin_x, footer_y - 4),
                color=GREEN, width=0.8
            )

            # Sahifa raqami
            page_num = f"– {page.number + 1} / {doc.page_count} –"
            page.insert_text(
                (w / 2 - 20, footer_y + 8),
                page_num, fontsize=8,
                color=DARK_BLUE, fontname="Helvetica-Bold"
            )

            # Manzil
            page.insert_text(
                (margin_x, footer_y + 18),
                footer_text, fontsize=6.5,
                color=GRAY, fontname="Helvetica"
            )

        # Yangi PDF bytes
        out = io.BytesIO()
        doc.save(out)
        doc.close()
        return out.getvalue()

    except Exception as e:
        # Xatolik bo'lsa asl bytes qaytaradi
        return pdf_bytes


def _get_pdf_bytes(article_file):
    """
    FileField dan PDF bytes oladi.
    Lokal: disk dan, Cloudinary/URL: HTTP orqali.
    """
    import os
    # Lokal
    try:
        path = article_file.path
        if os.path.exists(path):
            with open(path, 'rb') as f:
                return f.read()
    except Exception:
        pass
    # URL (Cloudinary)
    try:
        import requests as req
        resp = req.get(article_file.url, timeout=30)
        if resp.status_code == 200:
            return resp.content
    except Exception:
        pass
    return None


def download_pdf(request, pk):
    """
    Word (.docx) faylni PDF ga aylantirib qaytaradi.
    Kolontitullar bilan birga xhtml2pdf shablon ishlatiladi.
    """
    from django.http import HttpResponse, Http404
    from django.shortcuts import redirect as _redirect
    from django.template.loader import render_to_string

    article = get_object_or_404(Article, pk=pk, status='published')
    Article.objects.filter(pk=pk).update(downloads_count=article.downloads_count + 1)

    safe_title = article.title[:40].replace(' ', '_').replace('/', '_').replace('\\', '_')
    filename = f"TIFT_{safe_title}.pdf"

    if not article.pdf_file:
        raise Http404("Maqola fayli topilmadi.")

    try:
        from xhtml2pdf import pisa
        import io

        content_html = _get_article_content_html(article.pdf_file)
        html_string = render_to_string('article_pdf.html', {
            'article': article,
            'request': request,
            'pdf_content_html': content_html,
        })
        buffer = io.BytesIO()
        res = pisa.CreatePDF(src=html_string, dest=buffer, encoding='utf-8')
        if not res.err:
            response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
    except Exception:
        pass

    # Fallback: asl faylni qaytarish
    return _redirect(article.pdf_file.url)


def generate_article_pdf(request, pk):
    """Maqolani brauzerda ko'rish uchun PDF ga aylantiradi (Word → PDF)."""
    from django.template.loader import render_to_string
    from django.http import HttpResponse
    import io

    try:
        from xhtml2pdf import pisa
    except ImportError:
        return HttpResponse("xhtml2pdf o'rnatilmagan.", status=500)

    article = get_object_or_404(Article, pk=pk, status='published')
    if not request.session.get(f'pdf_viewed_{pk}'):
        Article.objects.filter(pk=pk).update(downloads_count=article.downloads_count + 1)
        request.session[f'pdf_viewed_{pk}'] = True

    content_html = ''
    if article.pdf_file:
        try:
            content_html = _get_article_content_html(article.pdf_file)
        except Exception:
            pass

    html_string = render_to_string('article_pdf.html', {
        'article': article,
        'request': request,
        'pdf_content_html': content_html,
    })
    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(src=html_string, dest=buffer, encoding='utf-8')

    if pisa_status.err:
        return HttpResponse("PDF yaratishda xatolik yuz berdi.", status=500)

    safe_title = article.title[:40].replace(' ', '_').replace('/', '_').replace('\\', '_')
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="TIFT_{safe_title}.pdf"'
    return response


def signup(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('my_articles')
    else:
        form = CustomUserCreationForm()
    return render(request, 'registration/signup.html', {'form': form})
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('my_articles')
    else:
        form = CustomUserCreationForm()
    return render(request, 'registration/signup.html', {'form': form})


@login_required
def submit_article(request):
    if request.method == 'POST':
        form = ArticleSubmissionForm(request.POST, request.FILES)
        if form.is_valid():
            article = form.save(commit=False)
            article.author = request.user
            article.status = 'submitted'
            article.save()
            return redirect('my_articles')
    else:
        form = ArticleSubmissionForm()
    return render(request, 'submit_article.html', {'form': form})


@login_required
def my_articles(request):
    articles = Article.objects.filter(author=request.user).order_by('-created_at')
    status_counts = {
        'submitted':    articles.filter(status='submitted').count(),
        'under_review': articles.filter(status='under_review').count(),
        'accepted':     articles.filter(status='accepted').count(),
        'published':    articles.filter(status='published').count(),
        'rejected':     articles.filter(status='rejected').count(),
    }
    return render(request, 'my_articles.html', {
        'articles': articles,
        'status_counts': status_counts,
    })


@login_required
def profile(request):
    """Profil ko'rish + ma'lumotlarni yangilash"""
    if request.method == 'POST':
        form = ProfileUpdateForm(request.POST, request.FILES, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Profil muvaffaqiyatli yangilandi!')
            return redirect('profile')
        else:
            messages.error(request, 'Xatolik yuz berdi. Iltimos, qaytadan urinib ko\'ring.')
    else:
        form = ProfileUpdateForm(instance=request.user)

    articles = Article.objects.filter(author=request.user).order_by('-created_at')
    status_counts = {
        'submitted':    articles.filter(status='submitted').count(),
        'under_review': articles.filter(status='under_review').count(),
        'accepted':     articles.filter(status='accepted').count(),
        'published':    articles.filter(status='published').count(),
        'rejected':     articles.filter(status='rejected').count(),
    }
    return render(request, 'profile.html', {
        'form': form,
        'articles': articles,
        'status_counts': status_counts,
    })


@login_required
def change_password(request):
    """Parol o'zgartirish"""
    if request.method == 'POST':
        form = CustomPasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)
            messages.success(request, 'Parol muvaffaqiyatli o\'zgartirildi!')
            return redirect('profile')
        else:
            messages.error(request, 'Xatolik yuz berdi.')
    else:
        form = CustomPasswordChangeForm(request.user)
    return render(request, 'change_password.html', {'form': form})


def conferences(request):
    items = Conference.objects.filter(is_active=True).order_by('-date')
    return render(request, 'conferences.html', {'items': items})


def news_list(request):
    items = News.objects.filter(is_active=True).order_by('-created_at')
    return render(request, 'news.html', {'items': items})


def documents(request):
    normative = Document.objects.filter(is_active=True, category='normative').order_by('order')
    requirements = Document.objects.filter(is_active=True, category='requirement').order_by('order')
    templates = Document.objects.filter(is_active=True, category='template').order_by('order')
    other = Document.objects.filter(is_active=True, category='other').order_by('order')
    return render(request, 'documents.html', {
        'normative': normative,
        'requirements': requirements,
        'templates': templates,
        'other': other,
    })
